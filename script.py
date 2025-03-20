import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import RGBColor
from datetime import datetime
from docx.text.run import *
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P

def add_cover_page(doc, title):
    # Add cover page
    doc.add_paragraph().add_run().add_break()  # Add some space at top
    
    # Add title
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.font.size = Pt(24)
    run.font.bold = True
    
    # Add date
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(12)
    
    # Add page break
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def add_table_of_contents(doc):
    # Add TOC title
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Table of Contents")
    run.font.size = Pt(16)
    run.font.bold = True
    
    # Add placeholder for TOC
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Right-click here to update field.")
    
    # Add page break
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def add_page_numbers(doc):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add "Page X of Y" format
    run = paragraph.add_run('Page ')
    
    # Add page number (X)
    run = paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fld_char1)
    
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'PAGE'
    run._element.append(instr_text)
    
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._element.append(fld_char2)
    

def post_process(doc, filename):
    try:
        # Load the cover page template
        cover_doc = Document('cover page.docx')
        
        # Create a new document for combining
        combined_doc = Document()
        
        # Get policy name from filename
        policy_name = filename.replace('.md', '').replace('_', ' ').title()
        policy_version = "1.0.0"
        
        # Copy cover page content and replace placeholder "Policy Name & Version", f"{policy_name} {policy_version}"
        # First, replace the placeholder text in the cover page document
        for paragraph in cover_doc.paragraphs:
            if "Policy Name & Version" in paragraph.text:
                # Clear the paragraph
                p = paragraph._element
                p.clear_content()
                
                # Add new run with text and set its format
                new_run = paragraph.add_run(f"{policy_name} {policy_version}")
                new_run.font.size = Pt(22)  # 设置字体大小
                new_run.font.bold = True     # 如果需要粗体
                new_run.font.name = "Times New Roman" # 如果需要设置字体
                # new_run.font.color.rgb = RGBColor(0, 0, 0)  # 如果需要设置颜色
        
        # Now copy all elements to the combined document
        for element in cover_doc.element.body:
            combined_doc.element.body.append(element)
        
        # Copy main document content
        for element in doc.element.body:
            combined_doc.element.body.append(element)
        
        # Add page numbers
        add_page_numbers(combined_doc)
        
        return combined_doc
    except Exception as e:
        print(f"Error combining with cover page: {str(e)}")
        return doc

def convert_md_to_docx():
    # Get all markdown files in the markdown directory
    markdown_dir = 'markdown'
    for filename in os.listdir(markdown_dir):
        # Process all files and directories recursively
        for root, dirs, files in os.walk(markdown_dir):
            # Create corresponding directories in output
            rel_path = os.path.relpath(root, markdown_dir)
            docx_dir = os.path.join('', rel_path)
            os.makedirs(docx_dir, exist_ok=True)
            
            # Process markdown files
            for filename in files:
                if filename.endswith('.md'):
                    md_path = os.path.join(root, filename)
                    docx_path = os.path.join(docx_dir, filename.replace('.md', '.docx'))
                    
                    # Create a new Word document
                    doc = Document()
                    
                    # Get document title from filename
                    doc_title = filename.replace('.md', '').replace('_', ' ').title()
                    
                    # Add cover page
                    # add_cover_page(doc, doc_title)
                    
                    # Add table of contents
                    # add_table_of_contents(doc)
                    
                    # Read and process the markdown file
                    with open(md_path, 'r', encoding='utf-8') as md_file:
                        lines = md_file.readlines()
                        
                        for line in lines:
                            # Handle headers
                            if line.startswith('#'):
                                level = line.count('#')
                                text = line.strip('#').strip()
                                paragraph = doc.add_paragraph()
                                run = paragraph.add_run(text)
                                run.font.size = Pt(16 - level)  # Decrease size for deeper headers
                                run.bold = True
                                # Add header style for TOC
                                paragraph.style = f'Heading {level}'
                                
                            # Handle bullet points
                            elif line.strip().startswith('* ') and not line.strip().startswith('**'):
                                # Remove only the leading bullet point marker
                                text = line[line.find('* ') + 2:].strip()
                                paragraph = doc.add_paragraph(style='List Bullet')
                                
                                # Process bold and italic text in bullet points
                                while '**' in text or '*' in text:
                                    # Handle bold text first (**text**)
                                    bold_start = text.find('**')
                                    if bold_start != -1:
                                        bold_end = text.find('**', bold_start + 2)
                                        if bold_end != -1:
                                            # Add text before bold
                                            if bold_start > 0:
                                                paragraph.add_run(text[:bold_start])
                                            # Add bold text
                                            bold_text = text[bold_start + 2:bold_end]
                                            run = paragraph.add_run(bold_text)
                                            run.bold = True
                                            # Continue with remaining text
                                            text = text[bold_end + 2:]
                                            continue
                                    
                                    # Handle italic text (*text*)
                                    italic_start = text.find('*')
                                    if italic_start != -1:
                                        italic_end = text.find('*', italic_start + 1)
                                        if italic_end != -1:
                                            # Add text before italic
                                            if italic_start > 0:
                                                paragraph.add_run(text[:italic_start])
                                            # Add italic text
                                            italic_text = text[italic_start + 1:italic_end]
                                            run = paragraph.add_run(italic_text)
                                            run.italic = True
                                            # Continue with remaining text
                                            text = text[italic_end + 1:]
                                            continue
                                    
                                    break  # Break if no more formatting found
                                
                                # Add any remaining text
                                if text:
                                    paragraph.add_run(text)
                                
                            # Handle tables
                            elif '|' in line:
                                # Skip table formatting lines
                                if ':---' in line:
                                    continue
                                cells = [cell.strip() for cell in line.strip('|').split('|')]
                                if len(cells) > 1:
                                    table = doc.add_table(rows=1, cols=len(cells))
                                    for i, cell in enumerate(cells):
                                        table.cell(0, i).text = cell
                                        
                            # Regular text
                            else:
                                if line.strip():
                                    paragraph = doc.add_paragraph()
                                    text = line.strip()
                                    
                                    # Process bold and italic text
                                    while '**' in text or '*' in text:
                                        # Handle bold text first (**text**)
                                        bold_start = text.find('**')
                                        if bold_start != -1:
                                            bold_end = text.find('**', bold_start + 2)
                                            if bold_end != -1:
                                                # Add text before bold
                                                if bold_start > 0:
                                                    paragraph.add_run(text[:bold_start])
                                                # Add bold text
                                                bold_text = text[bold_start + 2:bold_end]
                                                run = paragraph.add_run(bold_text)
                                                run.bold = True
                                                # Continue with remaining text
                                                text = text[bold_end + 2:]
                                                continue
                                        
                                        # Handle italic text (*text*)
                                        italic_start = text.find('*')
                                        if italic_start != -1:
                                            italic_end = text.find('*', italic_start + 1)
                                            if italic_end != -1:
                                                # Add text before italic
                                                if italic_start > 0:
                                                    paragraph.add_run(text[:italic_start])
                                                # Add italic text
                                                italic_text = text[italic_start + 1:italic_end]
                                                run = paragraph.add_run(italic_text)
                                                run.italic = True
                                                # Continue with remaining text
                                                text = text[italic_end + 1:]
                                                continue
                                        
                                        break  # Break if no more formatting found
                                    
                                    # Add any remaining text
                                    if text:
                                        paragraph.add_run(text)
                    
                    # Combine with cover page
                    final_doc = post_process(doc, filename)
                    
                    # Save the Word document
                    final_doc.save(docx_path)
                    print(f'Converted {md_path} to Word format')
if __name__ == '__main__':
    convert_md_to_docx()
